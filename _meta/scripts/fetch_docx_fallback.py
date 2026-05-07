#!/usr/bin/env python3
"""
fetch_docx_fallback: 用 python-docx 抓 audit_2026-05-06 Phase 1 抓失败的 .docx 文件

来源:_meta/audit_2026-05-06/fetch_top600.log + fetch_retry.log 中
ok=False 且 url endswith .docx 的记录(去重)。

输出:
- staging: _meta/audit_2026-05-06/docx_staging/<hash>.md(参 normalize 流程做后续处理)
- 报告: _meta/audit_2026-05-06/docx_fetch_report.md

不直接入 vault — 入 vault 需要 trigger A 全套(rel_judge + 5C),
那需要 LLM 调用,放后续作业。本脚本只把 body 抓回来落 staging。
"""
from __future__ import annotations
import hashlib
import io
import json
import re
import sys
from datetime import datetime
from pathlib import Path

import requests
from docx import Document  # python-docx

VAULT = Path(__file__).resolve().parent.parent.parent
LOG_DIR = VAULT / "_meta" / "audit_2026-05-06"
STAGING = LOG_DIR / "docx_staging"
STAGING.mkdir(parents=True, exist_ok=True)
REPORT_MD = LOG_DIR / "docx_fetch_report.md"

NOW_ISO = datetime.now().strftime("%Y-%m-%dT%H:%M:%S+08:00")
TIMEOUT = 30

USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0 Safari/537.36"
)


def collect_docx_urls() -> list[str]:
    """从两个 log 收集 ok=False + docx 的 URL,去重"""
    urls = set()
    for log in [LOG_DIR / "fetch_top600.log", LOG_DIR / "fetch_retry.log"]:
        if not log.exists():
            continue
        for ln in log.read_text(encoding="utf-8", errors="ignore").splitlines():
            ln = ln.strip()
            if not ln:
                continue
            try:
                r = json.loads(ln)
            except json.JSONDecodeError:
                continue
            if r.get("ok") is False and isinstance(r.get("url"), str) and r["url"].lower().endswith(".docx"):
                urls.add(r["url"])
    return sorted(urls)


def fetch_docx(url: str) -> tuple[bool, str | None, str]:
    """download + parse docx → (ok, body_text, error_msg)"""
    try:
        resp = requests.get(url, timeout=TIMEOUT, headers={"User-Agent": USER_AGENT})
        resp.raise_for_status()
        if not resp.content:
            return False, None, "empty body"
        doc = Document(io.BytesIO(resp.content))
        paras = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                paras.append(t)
        # 表格内容也抓
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    paras.append(" | ".join(cells))
        body = "\n\n".join(paras)
        if len(body) < 50:
            return False, None, f"body too short ({len(body)} chars)"
        return True, body, ""
    except requests.RequestException as e:
        return False, None, f"http error: {str(e)[:120]}"
    except Exception as e:
        return False, None, f"parse error: {str(e)[:120]}"


def url_hash(url: str) -> str:
    return hashlib.sha256(url.encode("utf-8")).hexdigest()[:8]


def main() -> int:
    urls = collect_docx_urls()
    print(f"待抓 docx URLs(去重): {len(urls)}")
    if not urls:
        print("没找到 .docx 失败记录")
        return 0

    results = []
    ok_count = 0
    err_count = 0
    for i, url in enumerate(urls, 1):
        h = url_hash(url)
        print(f"[{i}/{len(urls)}] {h}: {url[:80]}")
        ok, body, err = fetch_docx(url)
        if ok:
            ok_count += 1
            # 落 staging
            staging_path = STAGING / f"{h}.md"
            fm = {
                "slug": h,
                "url": url,
                "fetched_at": NOW_ISO,
                "fetched_method": "python-docx",
                "content_type": "docx",
                "body_len": len(body),
            }
            fm_yaml = "\n".join(f"{k}: {json.dumps(v, ensure_ascii=False) if isinstance(v, str) else v}" for k, v in fm.items())
            staging_path.write_text(f"---\n{fm_yaml}\n---\n\n# {h}\n\n## 政策原文\n\n{body}\n", encoding="utf-8")
            print(f"  ✓ {len(body)} chars → {staging_path.name}")
        else:
            err_count += 1
            print(f"  ✗ {err}")
        results.append({"url": url, "hash": h, "ok": ok, "body_len": len(body) if body else 0, "error": err})

    # 写报告
    lines = [
        f"# DOCX 兜底抓取报告 — {NOW_ISO}",
        "",
        f"- 待抓 URL: **{len(urls)}**(去重 from fetch_top600.log + fetch_retry.log)",
        f"- 成功: **{ok_count}**",
        f"- 失败: **{err_count}**",
        f"- 成功率: {ok_count/len(urls)*100:.1f}%",
        "",
        "## 成功列表",
        "",
        "| hash | body_len | url |",
        "|---|---:|---|",
    ]
    for r in results:
        if r["ok"]:
            lines.append(f"| {r['hash']} | {r['body_len']} | {r['url'][:100]} |")
    lines += ["", "## 失败列表", "", "| hash | error | url |", "|---|---|---|"]
    for r in results:
        if not r["ok"]:
            lines.append(f"| {r['hash']} | {(r.get('error') or '')[:60]} | {r['url'][:100]} |")
    lines += [
        "",
        "## 下一步",
        "",
        "成功 staging 在 `_meta/audit_2026-05-06/docx_staging/<hash>.md`,",
        "走 normalize_to_raw.py 入 vault → trigger A 全套(rel_judge + 5C)",
        "需要 LLM 调用,作为后续作业。",
    ]
    REPORT_MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"\n报告 → {REPORT_MD}")
    print(f"staging → {STAGING}/")
    print(f"成功 {ok_count} / 失败 {err_count}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
